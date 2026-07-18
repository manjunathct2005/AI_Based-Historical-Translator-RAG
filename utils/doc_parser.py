"""
utils/doc_parser.py

Parsing for user-uploaded documents: PDF (text + scanned/image pages) and
Word (.docx). This is used only to pull text OUT of a file the user gives
us (e.g. a photographed inscription in a PDF, or a translation draft in
Word) -- it is NOT a local knowledge base. The RAG context itself always
comes from live web retrieval (see rag/ and web/).
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from typing import List

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    page_number: int
    text: str
    has_images: bool


@dataclass
class ParsedDocument:
    pages: List[ParsedPage]
    full_text: str
    source_type: str  # "pdf" | "docx"


def parse_pdf(file_bytes: bytes, extract_images_for_ocr: bool = True) -> ParsedDocument:
    """
    Extract text from a PDF using PyMuPDF (fast, good for embedded text)
    with a pdfplumber fallback for tricky layouts/tables. Pages that come
    back with near-zero text are flagged as likely scanned images so the
    caller can route them through OCR (see utils/ocr.py).
    """
    import fitz  # PyMuPDF

    pages: List[ParsedPage] = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            has_images = len(page.get_images(full=True)) > 0

            if len(text) < 20:
                # Likely a scanned page -- try pdfplumber as a second
                # opinion before giving up and telling the caller to OCR it.
                text = _pdfplumber_page_text(file_bytes, i) or text

            pages.append(ParsedPage(page_number=i + 1, text=text, has_images=has_images))
    finally:
        doc.close()

    full_text = "\n\n".join(p.text for p in pages if p.text)
    return ParsedDocument(pages=pages, full_text=full_text, source_type="pdf")


def _pdfplumber_page_text(file_bytes: bytes, page_index: int) -> str:
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if page_index < len(pdf.pages):
                return (pdf.pages[page_index].extract_text() or "").strip()
    except Exception:  # pragma: no cover - defensive
        logger.exception("pdfplumber fallback failed on page %s", page_index)
    return ""


def render_pdf_page_as_image(file_bytes: bytes, page_index: int, dpi: int = 300) -> bytes:
    """Render a PDF page to PNG bytes -- used to hand scanned pages to OCR."""
    import fitz

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        page = doc[page_index]
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix)
        return pix.tobytes("png")
    finally:
        doc.close()


def parse_docx(file_bytes: bytes) -> ParsedDocument:
    """Extract paragraphs (and simple table text) from a .docx file."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: List[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    full_text = "\n".join(parts)
    # docx has no fixed "pages" concept pre-render, so we treat it as one page.
    return ParsedDocument(
        pages=[ParsedPage(page_number=1, text=full_text, has_images=len(doc.inline_shapes) > 0)],
        full_text=full_text,
        source_type="docx",
    )


def write_docx(text: str, title: str = "Translation Output") -> bytes:
    """Produce a simple .docx from generated text, for download in the UI."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for para in text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
